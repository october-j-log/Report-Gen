from docx import Document

document = Document()

# Title
document.add_heading(
    'Rectifier Validation Report',
    level=1
)

# Objective
document.add_heading(
    '1. Objective',
    level=2
)

document.add_paragraph(
    'To evaluate the performance, efficiency and thermal behavior of the rectifier under specified operating conditions.'
)

# Test Setup
document.add_heading(
    '2. Test Setup',
    level=2
)

document.add_paragraph(
    'The rectifier DUT was tested using the validation test setup. Electrical parameters such as voltage, current, power and temperature were monitored and recorded throughout the test duration.'
)

# Validation Summary
document.add_heading(
    '3. Validation Summary',
    level=2
)

document.add_paragraph(
    'Validation results will be inserted here automatically.'
)

# Test Evaluation
document.add_heading(
    '4. Test Evaluation',
    level=2
)

document.add_paragraph(
    'PASS/FAIL evaluation will be inserted here automatically.'
)

# Observations
document.add_heading(
    '5. Observations',
    level=2
)

document.add_paragraph(
    'Test observations will be inserted here automatically.'
)

# Graphs
document.add_heading(
    '6. Graphs',
    level=2
)

document.add_paragraph(
    'Generated graphs will be inserted here automatically.'
)

# Conclusion
document.add_heading(
    '7. Conclusion',
    level=2
)

document.add_paragraph(
    'The conclusion of the validation test will be inserted here automatically.'
)

# Save Report
document.save(
    'reports/rectifier_report_v2.docx'
)

print("Report Generated Successfully!")