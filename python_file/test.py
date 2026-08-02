import pandas as pd
import matplotlib.pyplot as plt
import os
import re
from docx import Document
from docx.shared import Inches
import yaml

# -----------------------------
# LOAD CONFIGURATION
# -----------------------------
with open("config.yaml", "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

TH_EFFICIENCY = config["thresholds"]["efficiency_min"]
TH_TEMPERATURE = config["thresholds"]["temperature_max"]


# -----------------------------
# LOG FILE PARSER（支持文本日志）
# -----------------------------
def parse_log_file(file_path):
    """
    解析文本日志文件，提取电压、电流、温度等参数
    """
    data = {
        "rb_voltage": [],
        "rb_current": [],
        "Total_kW": [],
        "IVT_temp": []
    }

    patterns = [
        {
            "voltage": r'电压[=:]\s*([\d.]+)\s*V',
            "current": r'电流[=:]\s*([\d.]+)\s*A',
            "temp": r'温度[=:]\s*([\d.]+)\s*°?C'
        },
        {
            "voltage": r'[Vv]oltage\s*[:=]\s*([\d.]+)\s*V',
            "current": r'[Cc]urrent\s*[:=]\s*([\d.]+)\s*A',
            "temp": r'[Tt]emp(?:erature)?\s*[:=]\s*([\d.]+)\s*°?C'
        },
        {
            "voltage": r'[Vv]\s*=\s*([\d.]+)',
            "current": r'[Ii]\s*=\s*([\d.]+)',
            "temp": r'[Tt]\s*=\s*([\d.]+)'
        }
    ]

    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        if not line.strip() or line.strip().startswith('#'):
            continue

        matched = False
        for pattern in patterns:
            voltage_match = re.search(pattern["voltage"], line)
            current_match = re.search(pattern["current"], line)
            temp_match = re.search(pattern["temp"], line)

            if voltage_match and current_match:
                v = float(voltage_match.group(1))
                i = float(current_match.group(1))
                t = float(temp_match.group(1)) if temp_match else 25.0

                data["rb_voltage"].append(v)
                data["rb_current"].append(i)
                data["IVT_temp"].append(t)

                power_kw = (v * i) / 1000
                data["Total_kW"].append(power_kw / 0.95)

                matched = True
                break

        if not matched:
            numbers = re.findall(r'[\d.]+', line)
            if len(numbers) >= 3:
                try:
                    v = float(numbers[0])
                    i = float(numbers[1])
                    t = float(numbers[2]) if len(numbers) > 2 else 25.0

                    if 100 <= v <= 1000 and 1 <= i <= 100 and 0 <= t <= 100:
                        data["rb_voltage"].append(v)
                        data["rb_current"].append(i)
                        data["IVT_temp"].append(t)
                        power_kw = (v * i) / 1000
                        data["Total_kW"].append(power_kw / 0.95)
                except ValueError:
                    continue

    if len(data["rb_voltage"]) == 0:
        print("警告：未能从日志文件中提取到有效数据，请检查日志格式")
        return pd.DataFrame()

    return pd.DataFrame(data)


# -----------------------------
# SETUP PROJECT FOLDERS
# -----------------------------
print("Working Directory:", os.getcwd())

os.makedirs("graphs", exist_ok=True)
os.makedirs("reports", exist_ok=True)


# -----------------------------
# AUTO DETECT INPUT FILE
# -----------------------------
input_folder = "input"

SUPPORTED_EXTS = ['.csv', '.xlsx', '.xls', '.log', '.txt']

files = [f for f in os.listdir(input_folder)
         if any(f.endswith(ext) for ext in SUPPORTED_EXTS)]

if len(files) == 0:
    print(f"没有找到支持的文件，支持的格式: {', '.join(SUPPORTED_EXTS)}")
    exit()

file_path = os.path.join(input_folder, files[0])
print("Processing:", file_path)

if file_path.endswith('.csv'):
    df = pd.read_csv(file_path)
elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
    try:
        df = pd.read_excel(file_path, engine='openpyxl')
    except ImportError:
        print("请安装 openpyxl: pip install openpyxl")
        exit()
elif file_path.endswith('.log') or file_path.endswith('.txt'):
    df = parse_log_file(file_path)
    if df.empty:
        print("日志解析失败，请检查文件格式")
        exit()
else:
    print(f"不支持的文件格式: {file_path}")
    exit()

print(f"成功读取 {len(df)} 行数据")


# -----------------------------
# DATA PROCESSING
# -----------------------------
df["Output_Power_kW"] = (df["rb_voltage"] * df["rb_current"]) / 1000

df["Efficiency"] = df.apply(
    lambda x: (x["Output_Power_kW"] / x["Total_kW"] * 100)
    if x["Total_kW"] != 0 else None,
    axis=1
)

df_valid = df[df["rb_current"] > 0].copy()


# -----------------------------
# STATISTICS
# -----------------------------
max_eff = df_valid["Efficiency"].max()
min_eff = df_valid["Efficiency"].min()
avg_eff = df_valid["Efficiency"].mean()

max_power = df_valid["Output_Power_kW"].max()
avg_power = df_valid["Output_Power_kW"].mean()

max_temp = df_valid["IVT_temp"].max()
avg_temp = df_valid["IVT_temp"].mean()

avg_voltage = df_valid["rb_voltage"].mean()

# ========== 异常统计 ==========
total_samples = len(df_valid)

low_eff_count = (df_valid["Efficiency"] < TH_EFFICIENCY).sum()
low_eff_rate = low_eff_count / total_samples * 100 if total_samples > 0 else 0

high_temp_count = (df_valid["IVT_temp"] > TH_TEMPERATURE).sum()
high_temp_rate = high_temp_count / total_samples * 100 if total_samples > 0 else 0

low_eff_indices = df_valid.index[df_valid["Efficiency"] < TH_EFFICIENCY].tolist()
high_temp_indices = df_valid.index[df_valid["IVT_temp"] > TH_TEMPERATURE].tolist()


# -----------------------------
# PASS / FAIL LOGIC
# -----------------------------
efficiency_result = "PASS" if avg_eff >= TH_EFFICIENCY else "FAIL"
temperature_result = "PASS" if max_temp <= TH_TEMPERATURE else "FAIL"
overall_result = "PASS" if efficiency_result == "PASS" and temperature_result == "PASS" else "FAIL"


# -----------------------------
# PLOTS（带异常标记）
# -----------------------------
# 1. 效率图
plt.figure(figsize=(10, 5))
plt.plot(df_valid.index, df_valid["Efficiency"], 'b-', label='Efficiency', linewidth=1.5)

low_eff_mask = df_valid["Efficiency"] < TH_EFFICIENCY
if low_eff_mask.any():
    plt.scatter(
        df_valid.index[low_eff_mask],
        df_valid["Efficiency"][low_eff_mask],
        color='red', s=60, zorder=5, label=f'Below {TH_EFFICIENCY}%'
    )

plt.axhline(y=TH_EFFICIENCY, color='r', linestyle='--', linewidth=1.5, label=f'Threshold: {TH_EFFICIENCY}%')
plt.title("Efficiency vs Sample")
plt.xlabel("Sample")
plt.ylabel("Efficiency (%)")
plt.legend()
plt.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig("graphs/efficiency.png")
plt.close()

# 2. 输出电压图
plt.figure(figsize=(10, 5))
plt.plot(df_valid.index, df_valid["rb_voltage"], 'b-', linewidth=1.5)
plt.title("Output Voltage vs Sample")
plt.xlabel("Sample")
plt.ylabel("Voltage (V)")
plt.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig("graphs/output_voltage.png")
plt.close()

# 3. 温度图
plt.figure(figsize=(10, 5))
plt.plot(df_valid.index, df_valid["IVT_temp"], 'b-', label='Temperature', linewidth=1.5)

high_temp_mask = df_valid["IVT_temp"] > TH_TEMPERATURE
if high_temp_mask.any():
    plt.scatter(
        df_valid.index[high_temp_mask],
        df_valid["IVT_temp"][high_temp_mask],
        color='red', s=60, zorder=5, label=f'Above {TH_TEMPERATURE}°C'
    )

plt.axhline(y=TH_TEMPERATURE, color='r', linestyle='--', linewidth=1.5, label=f'Threshold: {TH_TEMPERATURE}°C')
plt.title("Temperature vs Sample")
plt.xlabel("Sample")
plt.ylabel("Temperature (°C)")
plt.legend()
plt.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig("graphs/temperature.png")
plt.close()

# 4. 输出功率图
plt.figure(figsize=(10, 5))
plt.plot(df_valid.index, df_valid["Output_Power_kW"], 'b-', linewidth=1.5)
plt.title("Output Power vs Sample")
plt.xlabel("Sample")
plt.ylabel("Power (kW)")
plt.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig("graphs/output_power.png")
plt.close()


# -----------------------------
# SAVE PROCESSED DATA
# -----------------------------
df.to_csv("csv file/processed_rectifier_data.csv", index=False)
print("\nProcessed data saved successfully!")


# -----------------------------
# PRINT SUMMARY（终端输出）
# -----------------------------
print("\n----- TEST EVALUATION -----")
print("Efficiency :", efficiency_result)
print("Temperature :", temperature_result)
print("Overall Result :", overall_result)

print("\n----- VALIDATION SUMMARY -----")
print(f"Maximum Efficiency : {max_eff:.2f}%")
print(f"Minimum Efficiency : {min_eff:.2f}%")
print(f"Average Efficiency : {avg_eff:.2f}%")
print(f"Maximum Output Power : {max_power:.2f} kW")
print(f"Average Output Power : {avg_power:.2f} kW")
print(f"Maximum Temperature : {max_temp:.2f} °C")
print(f"Average Temperature : {avg_temp:.2f} °C")
print(f"Average Output Voltage : {avg_voltage:.2f} V")

print("\n----- ANOMALY STATISTICS -----")
print(f"Total Samples: {total_samples}")
print(f"Samples with Efficiency < {TH_EFFICIENCY}%: {low_eff_count} ({low_eff_rate:.1f}%)")
print(f"Samples with Temperature > {TH_TEMPERATURE}°C: {high_temp_count} ({high_temp_rate:.1f}%)")
if low_eff_indices:
    print(f"Low efficiency at samples: {low_eff_indices[:10]}{'...' if len(low_eff_indices) > 10 else ''}")
if high_temp_indices:
    print(f"High temperature at samples: {high_temp_indices[:10]}{'...' if len(high_temp_indices) > 10 else ''}")


# -----------------------------
# REPORT FILE (validation_summary.txt)
# -----------------------------
with open("reports/validation_summary.txt", "w") as f:
    f.write("RECTIFIER VALIDATION SUMMARY\n")
    f.write("=" * 50 + "\n\n")

    f.write(f"Maximum Efficiency : {max_eff:.2f}%\n")
    f.write(f"Minimum Efficiency : {min_eff:.2f}%\n")
    f.write(f"Average Efficiency : {avg_eff:.2f}%\n\n")

    f.write(f"Maximum Output Power : {max_power:.2f} kW\n")
    f.write(f"Average Output Power : {avg_power:.2f} kW\n\n")

    f.write(f"Maximum Temperature : {max_temp:.2f} °C\n")
    f.write(f"Average Temperature : {avg_temp:.2f} °C\n\n")

    f.write(f"Average Output Voltage : {avg_voltage:.2f} V\n\n")

    f.write("TEST EVALUATION\n")
    f.write("=" * 50 + "\n")
    f.write(f"Efficiency : {efficiency_result}\n")
    f.write(f"Temperature : {temperature_result}\n")
    f.write(f"Overall Result : {overall_result}\n")

    f.write("\nANOMALY STATISTICS\n")
    f.write("=" * 50 + "\n")
    f.write(f"Total Samples: {total_samples}\n")
    f.write(f"Samples with Efficiency < {TH_EFFICIENCY}%: {low_eff_count} ({low_eff_rate:.1f}%)\n")
    f.write(f"Samples with Temperature > {TH_TEMPERATURE}°C: {high_temp_count} ({high_temp_rate:.1f}%)\n")

    if low_eff_indices:
        f.write(f"Low efficiency sample indices: {low_eff_indices}\n")
    if high_temp_indices:
        f.write(f"High temperature sample indices: {high_temp_indices}\n")

print("Validation Summary Generated Successfully")


# -----------------------------
# OBSERVATION FILE
# -----------------------------
with open("reports/observations.txt", "w") as f:
    f.write("OBSERVATIONS\n")
    f.write("=" * 50 + "\n\n")
    f.write(f"Maximum efficiency observed: {max_eff:.2f}%\n")
    f.write(f"Average efficiency observed: {avg_eff:.2f}%\n")
    f.write(f"Maximum output power: {max_power:.2f} kW\n")
    f.write(f"Maximum temperature: {max_temp:.2f} °C\n")

print("Observations Generated Successfully")


# -----------------------------
# WORD REPORT GENERATION
# -----------------------------
document = Document()

# Title
document.add_heading('Rectifier Validation Report', level=1)

# 1. Objective
document.add_heading('1. Objective', level=2)
document.add_paragraph(
    'To evaluate the performance, efficiency and thermal behavior of the rectifier under specified operating conditions.'
)

# 2. Test Setup
document.add_heading('2. Test Setup', level=2)
document.add_paragraph(
    'The rectifier DUT was tested using the validation test setup. Electrical parameters such as voltage, current, power and temperature were monitored and recorded throughout the test duration.'
)

# 3. Validation Summary
document.add_heading('3. Validation Summary', level=2)
document.add_paragraph(f"Maximum Efficiency : {max_eff:.2f}%")
document.add_paragraph(f"Minimum Efficiency : {min_eff:.2f}%")
document.add_paragraph(f"Average Efficiency : {avg_eff:.2f}%")
document.add_paragraph(f"Maximum Output Power : {max_power:.2f} kW")
document.add_paragraph(f"Average Output Power : {avg_power:.2f} kW")
document.add_paragraph(f"Maximum Temperature : {max_temp:.2f} °C")
document.add_paragraph(f"Average Temperature : {avg_temp:.2f} °C")
document.add_paragraph(f"Average Output Voltage : {avg_voltage:.2f} V")
document.add_paragraph(f"Total Samples: {total_samples}")
document.add_paragraph(
    f"Samples with Efficiency < {TH_EFFICIENCY}%: {low_eff_count} ({low_eff_rate:.1f}%)"
)
document.add_paragraph(
    f"Samples with Temperature > {TH_TEMPERATURE}°C: {high_temp_count} ({high_temp_rate:.1f}%)"
)

# 4. Test Evaluation
document.add_heading('4. Test Evaluation', level=2)
document.add_paragraph(f"Efficiency Result : {efficiency_result}")
document.add_paragraph(f"Temperature Result : {temperature_result}")
document.add_paragraph(f"Overall Result : {overall_result}")

# 5. Observations
document.add_heading('5. Observations', level=2)

if avg_eff >= TH_EFFICIENCY:
    document.add_paragraph(
        f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which meets the acceptance criteria (≥ {TH_EFFICIENCY}%)."
    )
else:
    document.add_paragraph(
        f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which is below the required {TH_EFFICIENCY}% acceptance criteria."
    )

if max_temp <= TH_TEMPERATURE:
    document.add_paragraph(
        f"Thermal performance was satisfactory with a maximum temperature of {max_temp:.2f} °C (≤ {TH_TEMPERATURE}°C)."
    )
else:
    document.add_paragraph(
        f"Thermal performance was unsatisfactory as the maximum temperature reached {max_temp:.2f} °C (> {TH_TEMPERATURE}°C)."
    )

document.add_paragraph(f"The maximum output power delivered during the test was {max_power:.2f} kW.")
document.add_paragraph(f"The average output voltage during the test was {avg_voltage:.2f} V.")

# 6. Graphs
document.add_page_break()
document.add_heading('6. Graphs', level=2)

efficiency_graph = os.path.abspath("graphs/efficiency.png")
voltage_graph = os.path.abspath("graphs/output_voltage.png")
temperature_graph = os.path.abspath("graphs/temperature.png")
power_graph = os.path.abspath("graphs/output_power.png")

print("Efficiency Graph Path:", efficiency_graph)
print("Voltage Graph Path:", voltage_graph)
print("Temperature Graph Path:", temperature_graph)
print("Power Graph Path:", power_graph)

document.add_paragraph("Efficiency Graph")
document.add_picture(efficiency_graph, width=Inches(6))

document.add_paragraph("Output Voltage Graph")
document.add_picture(voltage_graph, width=Inches(6))

document.add_paragraph("Temperature Graph")
document.add_picture(temperature_graph, width=Inches(6))

document.add_paragraph("Output Power Graph")
document.add_picture(power_graph, width=Inches(6))

# 7. Conclusion
document.add_heading('7. Conclusion', level=2)

if overall_result == "PASS":
    document.add_paragraph(
        f"The rectifier successfully passed the validation test. "
        f"The average efficiency was {avg_eff:.2f}% and the maximum temperature remained within acceptable limits. "
        f"The DUT is considered suitable for operation under the tested conditions."
    )
else:
    document.add_paragraph(
        f"The rectifier failed to meet all validation requirements. "
        f"The average efficiency was {avg_eff:.2f}% which resulted in a FAIL condition. "
        f"Further investigation and optimization are recommended."
    )

# Save Report
document.save("reports/rectifier_report.docx")
print("\nWord Report Generated Successfully!")

# import pandas as pd
# import matplotlib.pyplot as plt
# import os
# import re
# from docx import Document
# from docx.shared import Inches
# import yaml
#
# # -----------------------------
# # LOAD CONFIGURATION  优化部分（）
# # -----------------------------
# with open("config.yaml", "r", encoding="utf-8") as f:
#     config = yaml.safe_load(f)
#
# TH_EFFICIENCY = config["thresholds"]["efficiency_min"]
# TH_TEMPERATURE = config["thresholds"]["temperature_max"]
#
#
# # -----------------------------
# # LOG FILE PARSER（支持文本日志）（添加部分）
# # -----------------------------
#
# def parse_log_file(file_path):
#     """
#     解析文本日志文件，提取电压、电流、温度等参数
#     支持多种日志格式：
#       - 电压=680.2V, 电流=35.4A, 温度=36.5°C
#       - [INFO] Voltage: 680.2 V, Current: 35.4 A, Temp: 36.5 C
#       - 2026-08-01 14:23:15 | 电压 680.2V | 电流 35.4A | 温度 36.5°C
#     """
#     data = {
#         "rb_voltage": [],
#         "rb_current": [],
#         "Total_kW": [],
#         "IVT_temp": []
#     }
#
#     # 定义多种正则匹配模式（兼容不同日志格式）
#     patterns = [
#         # 格式1: 电压=680.2V, 电流=35.4A, 温度=36.5°C
#         {
#             "voltage": r'电压[=:]\s*([\d.]+)\s*V',
#             "current": r'电流[=:]\s*([\d.]+)\s*A',
#             "temp": r'温度[=:]\s*([\d.]+)\s*°?C'
#         },
#         # 格式2: Voltage: 680.2 V, Current: 35.4 A, Temp: 36.5 C
#         {
#             "voltage": r'[Vv]oltage\s*[:=]\s*([\d.]+)\s*V',
#             "current": r'[Cc]urrent\s*[:=]\s*([\d.]+)\s*A',
#             "temp": r'[Tt]emp(?:erature)?\s*[:=]\s*([\d.]+)\s*°?C'
#         },
#         # 格式3: V=680.2 I=35.4 T=36.5
#         {
#             "voltage": r'[Vv]\s*=\s*([\d.]+)',
#             "current": r'[Ii]\s*=\s*([\d.]+)',
#             "temp": r'[Tt]\s*=\s*([\d.]+)'
#         }
#     ]
#
#     with open(file_path, 'r', encoding='utf-8') as f:
#         lines = f.readlines()
#
#     for line in lines:
#         # 跳过空行和纯注释行
#         if not line.strip() or line.strip().startswith('#'):
#             continue
#
#         matched = False
#         for pattern in patterns:
#             voltage_match = re.search(pattern["voltage"], line)
#             current_match = re.search(pattern["current"], line)
#             temp_match = re.search(pattern["temp"], line)
#
#             if voltage_match and current_match:
#                 v = float(voltage_match.group(1))
#                 i = float(current_match.group(1))
#                 t = float(temp_match.group(1)) if temp_match else 25.0
#
#                 data["rb_voltage"].append(v)
#                 data["rb_current"].append(i)
#                 data["IVT_temp"].append(t)
#
#                 # 计算总功率（假设效率约 95%，可根据实际情况调整）
#                 # 如果日志中有功率值，可以优先使用日志中的值
#                 power_kw = (v * i) / 1000
#                 data["Total_kW"].append(power_kw / 0.95)
#
#                 matched = True
#                 break
#
#         # 如果所有模式都匹配不到，尝试通用数值提取
#         if not matched:
#             numbers = re.findall(r'[\d.]+', line)
#             if len(numbers) >= 3:
#                 try:
#                     v = float(numbers[0])
#                     i = float(numbers[1])
#                     t = float(numbers[2]) if len(numbers) > 2 else 25.0
#
#                     # 简单范围校验，避免误提取
#                     if 100 <= v <= 1000 and 1 <= i <= 100 and 0 <= t <= 100:
#                         data["rb_voltage"].append(v)
#                         data["rb_current"].append(i)
#                         data["IVT_temp"].append(t)
#                         power_kw = (v * i) / 1000
#                         data["Total_kW"].append(power_kw / 0.95)
#                 except ValueError:
#                     continue
#
#     if len(data["rb_voltage"]) == 0:
#         print("警告：未能从日志文件中提取到有效数据，请检查日志格式")
#         return pd.DataFrame()
#
#     return pd.DataFrame(data)
#
# # -----------------------------
# # SETUP PROJECT FOLDERS
# # -----------------------------
# print("Working Directory:", os.getcwd())
#
# os.makedirs("graphs", exist_ok=True)
# os.makedirs("reports", exist_ok=True)
#
# # # -----------------------------
# # # AUTO DETECT INPUT FILE
# # # -----------------------------
# # input_folder = "input"
# #
# # files = os.listdir(input_folder)
# #
# # if len(files) == 0:
# #     print("No file found in input folder")
# #     exit()
# #
# # file_path = os.path.join(input_folder, files[0])
# #
# # print("Processing:", file_path)
# #
# # df = pd.read_csv(file_path)
#
# # -----------------------------
# # AUTO DETECT INPUT FILE(修改部分)
# # -----------------------------
# input_folder = "input"
#
# # 支持的文件扩展名
# SUPPORTED_EXTS = ['.csv', '.xlsx', '.xls', '.log', '.txt']
#
# # 过滤出支持的文件
# files = [f for f in os.listdir(input_folder)
#          if any(f.endswith(ext) for ext in SUPPORTED_EXTS)]
#
# if len(files) == 0:
#     print(f"没有找到支持的文件，支持的格式: {', '.join(SUPPORTED_EXTS)}")
#     exit()
#
# file_path = os.path.join(input_folder, files[0])
# print("Processing:", file_path)
#
# # 根据文件扩展名选择读取方式
# if file_path.endswith('.csv'):
#     df = pd.read_csv(file_path)
# elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
#     try:
#         df = pd.read_excel(file_path, engine='openpyxl')
#     except ImportError:
#         print("请安装 openpyxl: pip install openpyxl")
#         exit()
# elif file_path.endswith('.log') or file_path.endswith('.txt'):
#     df = parse_log_file(file_path)
#     if df.empty:
#         print("日志解析失败，请检查文件格式")
#         exit()
# else:
#     print(f"不支持的文件格式: {file_path}")
#     exit()
#
# print(f"成功读取 {len(df)} 行数据")
#
# # file_path = os.path.join(input_folder, files[0])
# #
# # print("Processing:", file_path)
# #
# # # 根据文件扩展名选择读取方式
# # if file_path.endswith('.csv'):
# #     df = pd.read_csv(file_path)
# # elif file_path.endswith('.xlsx'):
# #     df = pd.read_excel(file_path, engine='openpyxl')
# # else:
# #     print("不支持的文件格式，请使用 .csv 或 .xlsx")
# #     exit()
#
# # -----------------------------
# # DATA PROCESSING
# # -----------------------------
# df["Output_Power_kW"] = (df["rb_voltage"] * df["rb_current"]) / 1000
#
# # Safe efficiency calculation
# df["Efficiency"] = df.apply(
#     lambda x: (x["Output_Power_kW"] / x["Total_kW"] * 100)
#     if x["Total_kW"] != 0 else None,
#     axis=1
# )
#
# # Remove OFF-state samples
# df_valid = df[df["rb_current"] > 0].copy()
#
# # -----------------------------
# # STATISTICS
# # -----------------------------
# max_eff = df_valid["Efficiency"].max()
# min_eff = df_valid["Efficiency"].min()
# avg_eff = df_valid["Efficiency"].mean()
#
# max_power = df_valid["Output_Power_kW"].max()
# avg_power = df_valid["Output_Power_kW"].mean()
#
# max_temp = df_valid["IVT_temp"].max()
# avg_temp = df_valid["IVT_temp"].mean()
#
# avg_voltage = df_valid["rb_voltage"].mean()
#
# # -----------------------------
# # STATISTICS（添加部分）
# # -----------------------------
# max_eff = df_valid["Efficiency"].max()
# min_eff = df_valid["Efficiency"].min()
# avg_eff = df_valid["Efficiency"].mean()
#
# max_power = df_valid["Output_Power_kW"].max()
# avg_power = df_valid["Output_Power_kW"].mean()
#
# max_temp = df_valid["IVT_temp"].max()
# avg_temp = df_valid["IVT_temp"].mean()
#
# avg_voltage = df_valid["rb_voltage"].mean()
#
# # ========== 新增：异常统计 ==========
# total_samples = len(df_valid)
#
# # 效率异常统计
# low_eff_count = (df_valid["Efficiency"] < TH_EFFICIENCY).sum()
# low_eff_rate = low_eff_count / total_samples * 100 if total_samples > 0 else 0
#
# # 温度异常统计
# high_temp_count = (df_valid["IVT_temp"] > TH_TEMPERATURE).sum()
# high_temp_rate = high_temp_count / total_samples * 100 if total_samples > 0 else 0
#
# # 找出具体的异常样本索引（用于详细报告）
# low_eff_indices = df_valid.index[df_valid["Efficiency"] < TH_EFFICIENCY].tolist()
# high_temp_indices = df_valid.index[df_valid["IVT_temp"] > TH_TEMPERATURE].tolist()
#
# # # -----------------------------
# # # PASS / FAIL LOGIC
# # # -----------------------------
# # efficiency_result = "PASS" if avg_eff >= 90 else "FAIL"
# # temperature_result = "PASS" if max_temp <= 50 else "FAIL"
# #
# # overall_result = "PASS" if efficiency_result == "PASS" and temperature_result == "PASS" else "FAIL"
#
# # -----------------------------
# # PASS / FAIL LOGIC（使用配置值）
# # -----------------------------
# efficiency_result = "PASS" if avg_eff >= TH_EFFICIENCY else "FAIL"
# temperature_result = "PASS" if max_temp <= TH_TEMPERATURE else "FAIL"
#
# # # -----------------------------
# # # PLOTS
# # # -----------------------------
# #
# # plt.figure(figsize=(10, 5))
# # plt.plot(df_valid["Efficiency"])
# # plt.title("Efficiency vs Sample")
# # plt.xlabel("Sample")
# # plt.ylabel("Efficiency (%)")
# # plt.grid(True)
# # plt.savefig("graphs/efficiency.png")
# # plt.close()
# #
# # plt.figure(figsize=(10, 5))
# # plt.plot(df_valid["rb_voltage"])
# # plt.title("Output Voltage vs Sample")
# # plt.xlabel("Sample")
# # plt.ylabel("Voltage (V)")
# # plt.grid(True)
# # plt.savefig("graphs/output_voltage.png")
# # plt.close()
# #
# # plt.figure(figsize=(10, 5))
# # plt.plot(df_valid["IVT_temp"])
# # plt.title("Temperature vs Sample")
# # plt.xlabel("Sample")
# # plt.ylabel("Temperature (°C)")
# # plt.grid(True)
# # plt.savefig("graphs/temperature.png")
# # plt.close()
# #
# # plt.figure(figsize=(10, 5))
# # plt.plot(df_valid["Output_Power_kW"])
# # plt.title("Output Power vs Sample")
# # plt.xlabel("Sample")
# # plt.ylabel("Power (kW)")
# # plt.grid(True)
# # plt.savefig("graphs/output_power.png")
# # plt.close()
#
# # -----------------------------
# # PLOTS（带异常标记）（修改部分）
# # -----------------------------
#
# # 1. 效率图 - 标记低于阈值的点
# plt.figure(figsize=(10, 5))
# plt.plot(df_valid.index, df_valid["Efficiency"], 'b-', label='Efficiency', linewidth=1.5)
#
# # 找出低于阈值的点
# low_eff_mask = df_valid["Efficiency"] < TH_EFFICIENCY
# if low_eff_mask.any():
#     plt.scatter(
#         df_valid.index[low_eff_mask],
#         df_valid["Efficiency"][low_eff_mask],
#         color='red', s=60, zorder=5, label=f'Below {TH_EFFICIENCY}%'
#     )
#
# plt.axhline(y=TH_EFFICIENCY, color='r', linestyle='--', linewidth=1.5, label=f'Threshold: {TH_EFFICIENCY}%')
# plt.title("Efficiency vs Sample")
# plt.xlabel("Sample")
# plt.ylabel("Efficiency (%)")
# plt.legend()
# plt.grid(True, alpha=0.3)
# plt.tight_layout()
# plt.savefig("graphs/efficiency.png")
# plt.close()
#
# # 2. 输出电压图
# plt.figure(figsize=(10, 5))
# plt.plot(df_valid.index, df_valid["rb_voltage"], 'b-', linewidth=1.5)
# plt.title("Output Voltage vs Sample")
# plt.xlabel("Sample")
# plt.ylabel("Voltage (V)")
# plt.grid(True, alpha=0.3)
# plt.tight_layout()
# plt.savefig("graphs/output_voltage.png")
# plt.close()
#
# # 3. 温度图 - 标记超过阈值的点
# plt.figure(figsize=(10, 5))
# plt.plot(df_valid.index, df_valid["IVT_temp"], 'b-', label='Temperature', linewidth=1.5)
#
# # 找出超过阈值的点
# high_temp_mask = df_valid["IVT_temp"] > TH_TEMPERATURE
# if high_temp_mask.any():
#     plt.scatter(
#         df_valid.index[high_temp_mask],
#         df_valid["IVT_temp"][high_temp_mask],
#         color='red', s=60, zorder=5, label=f'Above {TH_TEMPERATURE}°C'
#     )
#
# plt.axhline(y=TH_TEMPERATURE, color='r', linestyle='--', linewidth=1.5, label=f'Threshold: {TH_TEMPERATURE}°C')
# plt.title("Temperature vs Sample")
# plt.xlabel("Sample")
# plt.ylabel("Temperature (°C)")
# plt.legend()
# plt.grid(True, alpha=0.3)
# plt.tight_layout()
# plt.savefig("graphs/temperature.png")
# plt.close()
#
# # 4. 输出功率图
# plt.figure(figsize=(10, 5))
# plt.plot(df_valid.index, df_valid["Output_Power_kW"], 'b-', linewidth=1.5)
# plt.title("Output Power vs Sample")
# plt.xlabel("Sample")
# plt.ylabel("Power (kW)")
# plt.grid(True, alpha=0.3)
# plt.tight_layout()
# plt.savefig("graphs/output_power.png")
# plt.close()
#
# # -----------------------------
# # SAVE PROCESSED DATA
# # -----------------------------
# df.to_csv("csv file/processed_rectifier_data.csv", index=False)
#
# print("\nProcessed data saved successfully!")
#
# # -----------------------------
# # PRINT SUMMARY
# # -----------------------------
# print("\n----- TEST EVALUATION -----")
# print("Efficiency :", efficiency_result)
# print("Temperature :", temperature_result)
# print("Overall Result :", overall_result)
#
# print("\n----- VALIDATION SUMMARY -----")
# print("Maximum Efficiency :", round(max_eff, 2), "%")
# print("Minimum Efficiency :", round(min_eff, 2), "%")
# print("Average Efficiency :", round(avg_eff, 2), "%")
#
# print("Maximum Output Power :", round(max_power, 2), "kW")
# print("Average Output Power :", round(avg_power, 2), "kW")
#
# print("Maximum Temperature :", round(max_temp, 2), "°C")
# print("Average Temperature :", round(avg_temp, 2), "°C")
#
# print("Average Output Voltage :", round(avg_voltage, 2), "V")
#
# # -----------------------------
# # PRINT SUMMARY（添加部分）
# # -----------------------------
# print("\n----- TEST EVALUATION -----")
# print("Efficiency :", efficiency_result)
# print("Temperature :", temperature_result)
# print("Overall Result :", overall_result)
#
# print("\n----- VALIDATION SUMMARY -----")
# print("Maximum Efficiency :", round(max_eff, 2), "%")
# print("Minimum Efficiency :", round(min_eff, 2), "%")
# print("Average Efficiency :", round(avg_eff, 2), "%")
#
# print("Maximum Output Power :", round(max_power, 2), "kW")
# print("Average Output Power :", round(avg_power, 2), "kW")
#
# print("Maximum Temperature :", round(max_temp, 2), "°C")
# print("Average Temperature :", round(avg_temp, 2), "°C")
#
# print("Average Output Voltage :", round(avg_voltage, 2), "V")
#
# # ========== 新增：打印异常统计 ==========
# print("\n----- ANOMALY STATISTICS -----")
# print(f"Total Samples: {total_samples}")
# print(f"Samples with Efficiency < {TH_EFFICIENCY}%: {low_eff_count} ({low_eff_rate:.1f}%)")
# print(f"Samples with Temperature > {TH_TEMPERATURE}°C: {high_temp_count} ({high_temp_rate:.1f}%)")
# if low_eff_indices:
#     print(f"Low efficiency at samples: {low_eff_indices[:10]}{'...' if len(low_eff_indices) > 10 else ''}")
# if high_temp_indices:
#     print(f"High temperature at samples: {high_temp_indices[:10]}{'...' if len(high_temp_indices) > 10 else ''}")
#
# # -----------------------------
# # REPORT FILE (SUMMARY)
# # -----------------------------
# with open("reports/validation_summary.txt", "w") as f:
#
#     f.write("RECTIFIER VALIDATION SUMMARY\n")
#     f.write("=" * 50 + "\n\n")
#
#     f.write(f"Maximum Efficiency : {max_eff:.2f}%\n")
#     f.write(f"Minimum Efficiency : {min_eff:.2f}%\n")
#     f.write(f"Average Efficiency : {avg_eff:.2f}%\n\n")
#
#     f.write(f"Maximum Output Power : {max_power:.2f} kW\n")
#     f.write(f"Average Output Power : {avg_power:.2f} kW\n\n")
#
#     f.write(f"Maximum Temperature : {max_temp:.2f} °C\n")
#     f.write(f"Average Temperature : {avg_temp:.2f} °C\n\n")
#
#     f.write(f"Average Output Voltage : {avg_voltage:.2f} V\n\n")
#
#     f.write("TEST EVALUATION\n")
#     f.write("=" * 50 + "\n")
#     f.write(f"Efficiency : {efficiency_result}\n")
#     f.write(f"Temperature : {temperature_result}\n")
#     f.write(f"Overall Result : {overall_result}\n")
#
# print("Validation Summary Generated Successfully")
#
# # -----------------------------
# # REPORT FILE (SUMMARY)（添加部分）
# # -----------------------------
# with open("reports/validation_summary.txt", "w") as f:
#     f.write("RECTIFIER VALIDATION SUMMARY\n")
#     f.write("=" * 50 + "\n\n")
#
#     f.write(f"Maximum Efficiency : {max_eff:.2f}%\n")
#     f.write(f"Minimum Efficiency : {min_eff:.2f}%\n")
#     f.write(f"Average Efficiency : {avg_eff:.2f}%\n\n")
#
#     f.write(f"Maximum Output Power : {max_power:.2f} kW\n")
#     f.write(f"Average Output Power : {avg_power:.2f} kW\n\n")
#
#     f.write(f"Maximum Temperature : {max_temp:.2f} °C\n")
#     f.write(f"Average Temperature : {avg_temp:.2f} °C\n\n")
#
#     f.write(f"Average Output Voltage : {avg_voltage:.2f} V\n\n")
#
#     f.write("TEST EVALUATION\n")
#     f.write("=" * 50 + "\n")
#     f.write(f"Efficiency : {efficiency_result}\n")
#     f.write(f"Temperature : {temperature_result}\n")
#     f.write(f"Overall Result : {overall_result}\n")
#
#     # ========== 新增：异常统计 ==========
#     f.write("\nANOMALY STATISTICS\n")
#     f.write("=" * 50 + "\n")
#     f.write(f"Total Samples: {total_samples}\n")
#     f.write(f"Samples with Efficiency < {TH_EFFICIENCY}%: {low_eff_count} ({low_eff_rate:.1f}%)\n")
#     f.write(f"Samples with Temperature > {TH_TEMPERATURE}°C: {high_temp_count} ({high_temp_rate:.1f}%)\n")
#
#     if low_eff_indices:
#         f.write(f"Low efficiency sample indices: {low_eff_indices}\n")
#     if high_temp_indices:
#         f.write(f"High temperature sample indices: {high_temp_indices}\n")
#
#
# # -----------------------------
# # OBSERVATION FILE
# # -----------------------------
# with open("reports/observations.txt", "w") as f:
#
#     f.write("OBSERVATIONS\n")
#     f.write("=" * 50 + "\n\n")
#
#     f.write(f"Maximum efficiency observed: {max_eff:.2f}%\n")
#     f.write(f"Average efficiency observed: {avg_eff:.2f}%\n")
#     f.write(f"Maximum output power: {max_power:.2f} kW\n")
#     f.write(f"Maximum temperature: {max_temp:.2f} °C\n")
#
# print("Observations Generated Successfully")
# # -----------------------------
# # WORD REPORT GENERATION
# # -----------------------------
# document = Document()
#
# # Title
# document.add_heading(
#     'Rectifier Validation Report',
#     level=1
# )
#
# # Objective
# document.add_heading(
#     '1. Objective',
#     level=2
# )
#
# document.add_paragraph(
#     'To evaluate the performance, efficiency and thermal behavior of the rectifier under specified operating conditions.'
# )
#
# # Test Setup
# document.add_heading(
#     '2. Test Setup',
#     level=2
# )
#
# document.add_paragraph(
#     'The rectifier DUT was tested using the validation test setup. Electrical parameters such as voltage, current, power and temperature were monitored and recorded throughout the test duration.'
# )
#
# # Validation Summary
# document.add_heading(
#     '3. Validation Summary',
#     level=2
# )
#
# document.add_paragraph(
#     f"Maximum Efficiency : {max_eff:.2f}%"
# )
#
# document.add_paragraph(
#     f"Minimum Efficiency : {min_eff:.2f}%"
# )
#
# document.add_paragraph(
#     f"Average Efficiency : {avg_eff:.2f}%"
# )
#
# document.add_paragraph(
#     f"Maximum Output Power : {max_power:.2f} kW"
# )
#
# document.add_paragraph(
#     f"Average Output Power : {avg_power:.2f} kW"
# )
#
# document.add_paragraph(
#     f"Maximum Temperature : {max_temp:.2f} °C"
# )
#
# document.add_paragraph(
#     f"Average Temperature : {avg_temp:.2f} °C"
# )
#
# document.add_paragraph(
#     f"Average Output Voltage : {avg_voltage:.2f} V"
# )
#
# # Test Evaluation
# document.add_heading(
#     '4. Test Evaluation',
#     level=2
# )
#
# document.add_paragraph(
#     f"Efficiency Result : {efficiency_result}"
# )
#
# document.add_paragraph(
#     f"Temperature Result : {temperature_result}"
# )
#
# document.add_paragraph(
#     f"Overall Result : {overall_result}"
# )
#
# # Observations
# document.add_heading(
#     '5. Observations',
#     level=2
# )
#
# # if avg_eff >= 90:
# #     document.add_paragraph(
# #         f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which meets the acceptance criteria."
# #     )
# # else:
# #     document.add_paragraph(
# #         f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which is below the required 90% acceptance criteria."
# #     )
# #
# # if max_temp <= 50:
# #     document.add_paragraph(
# #         f"Thermal performance was satisfactory with a maximum temperature of {max_temp:.2f} °C."
# #     )
# # else:
# #     document.add_paragraph(
# #         f"Thermal performance was unsatisfactory as the maximum temperature reached {max_temp:.2f} °C."
# #     )
#
# if avg_eff >= TH_EFFICIENCY:
#     document.add_paragraph(
#         f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which meets the acceptance criteria (≥ {TH_EFFICIENCY}%)."
#     )
# else:
#     document.add_paragraph(
#         f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which is below the required {TH_EFFICIENCY}% acceptance criteria."
#     )
#
# if max_temp <= TH_TEMPERATURE:
#     document.add_paragraph(
#         f"Thermal performance was satisfactory with a maximum temperature of {max_temp:.2f} °C (≤ {TH_TEMPERATURE}°C)."
#     )
# else:
#     document.add_paragraph(
#         f"Thermal performance was unsatisfactory as the maximum temperature reached {max_temp:.2f} °C (> {TH_TEMPERATURE}°C)."
#     )
#
# document.add_paragraph(
#     f"The maximum output power delivered during the test was {max_power:.2f} kW."
# )
#
# document.add_paragraph(
#     f"The average output voltage during the test was {avg_voltage:.2f} V."
# )
# # Graphs
# document.add_page_break()
#
# document.add_heading(
#     '6. Graphs',
#     level=2
# )
#
# efficiency_graph = os.path.abspath("graphs/efficiency.png")
# voltage_graph = os.path.abspath("graphs/output_voltage.png")
# temperature_graph = os.path.abspath("graphs/temperature.png")
# power_graph = os.path.abspath("graphs/output_power.png")
#
# print("Efficiency Graph Path:", efficiency_graph)
# print("Voltage Graph Path:", voltage_graph)
# print("Temperature Graph Path:", temperature_graph)
# print("Power Graph Path:", power_graph)
#
# document.add_paragraph("Efficiency Graph")
# document.add_picture(
#     efficiency_graph,
#     width=Inches(6)
# )
#
# document.add_paragraph("Output Voltage Graph")
# document.add_picture(
#     voltage_graph,
#     width=Inches(6)
# )
#
# document.add_paragraph("Temperature Graph")
# document.add_picture(
#     temperature_graph,
#     width=Inches(6)
# )
#
# document.add_paragraph("Output Power Graph")
# document.add_picture(
#     power_graph,
#     width=Inches(6)
# )
# # Conclusion
# document.add_heading(
#     '7. Conclusion',
#     level=2
# )
#
# if overall_result == "PASS":
#
#     document.add_paragraph(
#         f"The rectifier successfully passed the validation test. "
#         f"The average efficiency was {avg_eff:.2f}% and the maximum temperature remained within acceptable limits. "
#         f"The DUT is considered suitable for operation under the tested conditions."
#     )
#
# else:
#
#     document.add_paragraph(
#         f"The rectifier failed to meet all validation requirements. "
#         f"The average efficiency was {avg_eff:.2f}% which resulted in a FAIL condition. "
#         f"Further investigation and optimization are recommended."
#     )
# # Save Report
# document.save(
#     "reports/rectifier_report.docx"
# )
#
# print("\nWord Report Generated Successfully!")